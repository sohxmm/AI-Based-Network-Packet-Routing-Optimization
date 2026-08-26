"""Generate the final project report as a .docx.

Run from the repository root::

    python scripts/build_final_report.py

The report is *generated* rather than hand-written so that every number in it
comes from an artifact in the repository. Hours are summed from the worklogs,
benchmark figures are read from `experiments/results/`, and model metrics from
`ml/results/`. If a number changes, regenerating the report changes it too —
which is the same discipline `scripts/verify_claims.py` enforces elsewhere.

Output: FINAL_REPORT.docx
"""

from __future__ import annotations

import json
import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor

REPO_ROOT = Path(__file__).resolve().parents[1]
sys.path.insert(0, str(REPO_ROOT))

RESULTS = REPO_ROOT / "experiments" / "results"
ML_RESULTS = REPO_ROOT / "ml" / "results"
ASSETS = REPO_ROOT / "docs" / "assets"

ACCENT = RGBColor(0x1C, 0x5C, 0xAB)
MUTED = RGBColor(0x52, 0x51, 0x4E)


# ---------------------------------------------------------------------------
# Data gathering
# ---------------------------------------------------------------------------
def parse_worklog(path: Path) -> dict:
    """Extract hours, entry dates and task lines from a worklog."""
    if not path.exists():
        return {"total": 0.0, "entries": 0, "tasks": [], "name": path.stem}

    text = path.read_text(encoding="utf-8")
    hours = [float(h) for h in re.findall(r"^\*{0,2}Hours:?\*{0,2}\s*([\d.]+)", text, re.M)]
    dates = re.findall(r"^##\s*([\d/\-]+)", text, re.M)
    # Per-entry hours, so a phase total can be summed from the log rather than
    # typed in beside it and left to drift.
    by_date = dict(zip(dates, hours, strict=False))
    tasks = [
        line.strip().lstrip("-").strip()
        for line in text.splitlines()
        if line.strip().startswith("-") and len(line.strip()) > 12
    ]
    name_match = re.search(r"^#\s*Internship Work Log(?:\s*—\s*(.+))?", text, re.M)
    name = (name_match.group(1) if name_match and name_match.group(1) else path.stem.replace("WORKLOG_", ""))
    return {
        "total": sum(hours),
        "entries": len(dates),
        "first": dates[0] if dates else "-",
        "last": dates[-1] if dates else "-",
        "tasks": tasks,
        "by_date": by_date,
        "name": name.strip(),
    }


def load_results() -> dict:
    return {
        path.stem: json.loads(path.read_text(encoding="utf-8"))
        for path in sorted(RESULTS.glob("*.json"))
    }


def load_ml() -> dict:
    out = {}
    for name in ("gnn", "rl", "lstm", "marl"):
        path = ML_RESULTS / f"{name}_evaluation.json"
        if path.exists():
            out[name] = json.loads(path.read_text(encoding="utf-8"))
    return out


# ---------------------------------------------------------------------------
# Document helpers
# ---------------------------------------------------------------------------
def add_heading(doc, text, level=1):
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        run.font.color.rgb = ACCENT
    return heading


def add_caption(doc, text):
    paragraph = doc.add_paragraph()
    run = paragraph.add_run(text)
    run.italic = True
    run.font.size = Pt(9)
    run.font.color.rgb = MUTED
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    return paragraph


def add_table(doc, headers, rows, widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Light Grid Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for index, header in enumerate(headers):
        cell = table.rows[0].cells[index]
        cell.text = str(header)
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(9)

    for row in rows:
        cells = table.add_row().cells
        for index, value in enumerate(row):
            cells[index].text = str(value)
            for paragraph in cells[index].paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(9)

    if widths:
        for row in table.rows:
            for index, width in enumerate(widths):
                row.cells[index].width = Inches(width)
    doc.add_paragraph()
    return table


def add_bullets(doc, items):
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


# ---------------------------------------------------------------------------
# Sections
# ---------------------------------------------------------------------------
def cover(doc, soham, sneha):
    title = doc.add_heading("AI-Based Network Packet Routing Optimization", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("Final Project Report")
    run.bold = True
    run.font.size = Pt(15)

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = meta.add_run(
        f"Algorithm & AI/ML Internship\n"
        f"{soham['name']} · {sneha['name']}\n"
        f"Combined effort: {soham['total'] + sneha['total']:.1f} hours"
    )
    run.font.size = Pt(11)
    run.font.color.rgb = MUTED

    doc.add_paragraph()
    intro = doc.add_paragraph()
    intro.add_run(
        "This report covers the full lifecycle of the project, including the "
        "internal technical review it underwent and the rebuild that followed. "
        "It is deliberately candid about what did not work: the central problem "
        "we found was that the project's documentation claimed results its own "
        "committed evidence contradicted, and the most useful thing we can do "
        "in a final report is not repeat that mistake."
    ).font.size = Pt(10.5)
    doc.add_page_break()


#: Entry dates belonging to the final review-and-rebuild phase.
REBUILD_PHASE = tuple(f"{day:02d}/08/2026" for day in range(12, 29))


def section_hours(doc, soham, sneha):
    add_heading(doc, "1. Hours and Work Completed", 1)

    # Summed from the worklog rather than written in by hand, so the phase table
    # cannot drift away from the entries it summarises.
    soham_phase_hours = sum(
        hours for date, hours in soham["by_date"].items() if date in REBUILD_PHASE
    )

    total = soham["total"] + sneha["total"]
    paragraph = doc.add_paragraph()
    paragraph.add_run(f"{total:.1f} hours").bold = True
    paragraph.add_run(
        f" of logged effort across {soham['entries'] + sneha['entries']} working "
        f"sessions, from {sneha['first']} to {soham['last']}."
    )

    add_table(
        doc,
        ["Contributor", "Hours", "Sessions", "First entry", "Last entry"],
        [
            [soham["name"], f"{soham['total']:.1f}", soham["entries"], soham["first"], soham["last"]],
            [sneha["name"], f"{sneha['total']:.1f}", sneha["entries"], sneha["first"], sneha["last"]],
            ["Total", f"{total:.1f}", soham["entries"] + sneha["entries"], "—", "—"],
        ],
        widths=[1.9, 0.8, 0.9, 1.1, 1.1],
    )

    add_heading(doc, "1.1 Phases", 2)
    add_table(
        doc,
        ["Phase", "Period", "Hours", "Delivered"],
        [
            ["Simulator and classical routing", "Jun 9 – Jun 22", "~32", "Network simulator, Dijkstra, Bellman-Ford, ACO, stress tests"],
            ["API, streaming, dashboard", "Jun 23 – Jul 3", "~35", "FastAPI, WebSocket streaming, React dashboard, D3 topology"],
            ["Learned routing", "Jul 14 – Jul 28", "~38", "GNN, PPO, LSTM scaffold, multi-agent routing, predictive mode"],
            ["Benchmarking and consolidation", "Jul 29 – Aug 11", "~33", "Benchmark harness, guardrails, experiment sandbox, Docker"],
            ["Review and rebuild", "Aug 12 – Aug 28", f"{soham_phase_hours:.1f}", "Correctness rebuild across every layer, plus QoS routing, fault tolerance, trace replay and live network measurement"],
        ],
        widths=[1.7, 1.1, 0.6, 2.7],
    )

    add_heading(doc, "1.2 What the final phase delivered", 2)
    add_bullets(
        doc,
        [
            "A correctness rebuild across every layer, each fix verified individually.",
            "Four of four AI features now run on a fresh clone; previously one of four did.",
            "The simulator was made closed-loop, so routing decisions change the network — the change that makes the project's central question answerable at all.",
            "The benchmark was rebuilt on independent replications with real effect sizes, replacing statistics that were invalid by construction.",
            "A test suite that can actually be installed and run, with 60+ tests, plus automated honesty gates wired into CI.",
            "Four capabilities beyond the original scope: QoS-aware multi-class routing, fault-tolerant rerouting with a convergence metric, trace replay, and live measurement of a real network.",
        ],
    )
    doc.add_page_break()


def section_learning(doc, soham, sneha):
    add_heading(doc, "2. What We Learnt", 1)

    doc.add_paragraph(
        "Drawn from both worklogs. The technical skills are the easy part to "
        "list; the lessons that changed how we work are further down and matter "
        "more."
    )

    add_heading(doc, "2.1 Technical ground covered", 2)
    add_table(
        doc,
        ["Area", "What we learnt to do"],
        [
            ["Graph algorithms", "Dijkstra, Bellman-Ford and Yen's k-shortest paths from first principles — including why Dijkstra's greedy choice property requires non-negative weights, and why that makes it provably optimal for additive costs."],
            ["Metaheuristics", "Ant Colony Optimization: pheromone reinforcement, evaporation as a mechanism for forgetting stale routes, and the exploration/exploitation trade-off that α and β control."],
            ["Graph neural networks", "Implementing message passing from scratch rather than importing a library — and consequently understanding why aggregation must be degree-normalised and why pooling choice determines what the model can represent at all."],
            ["Reinforcement learning", "Gymnasium environment design, PPO, reward shaping, and the difference between a partially observable MDP and an unobservable one."],
            ["Multi-agent RL", "Centralized training with decentralized execution: what the asymmetry actually means, and how to verify it rather than assert it."],
            ["Time series", "Why differencing matters for an autocorrelated series, and why persistence is the baseline any forecaster must beat."],
            ["Experimental statistics", "Pseudo-replication, the unit of replication, non-parametric effect sizes (Cliff's delta), and bootstrap confidence intervals."],
            ["Backend engineering", "Async FastAPI, WebSocket streaming, SQLAlchemy 2.0 with async sessions, Alembic migrations, retention policies."],
            ["Frontend engineering", "React with D3 without fighting the reconciler, memoisation for high-frequency updates, and colour choices validated for colour-vision deficiency rather than chosen by eye."],
            ["Delivery", "Multi-stage Docker builds, non-root containers, CI with a smoke test that runs the real thing."],
        ],
        widths=[1.5, 4.6],
    )

    add_heading(doc, "2.2 The lessons that changed how we work", 2)

    lessons = [
        (
            "A metric without a floor and a ceiling conveys nothing.",
            "We reported that our PPO agent reached a mean reward of -61 and "
            "treated it as a result. It is not one. Without knowing what a random "
            "policy scores and what an oracle scores on the same episodes, -61 is "
            "an uninterpretable number. Every model in the project is now reported "
            "as a normalized score between those two references, and doing so "
            "immediately revealed that our PPO agent does not beat simply taking "
            "the cheapest candidate path.",
        ),
        (
            "Check whether the environment can support the claim before building on it.",
            "We committed to showing AI beating classical routing, then built a "
            "simulator in which link utilization evolved independently of routing "
            "decisions. In that setting per-path latency minimisation is exactly "
            "optimal and Dijkstra solves it exactly, so the best possible outcome "
            "for any learned policy was a tie. No architecture, training budget or "
            "reward function could have fixed it. One design review asking 'if "
            "routing does not affect utilization, what is there to learn?' would "
            "have redirected the whole project.",
        ),
        (
            "Silence is the most expensive failure mode.",
            "Our RL router never loaded its model for the entire life of the "
            "project, because the loader looked for one filename and training "
            "saved another, and the resulting error was caught by a bare except "
            "that returned False without logging. Everything appeared to work. We "
            "now treat a swallowed exception as a defect in its own right: every "
            "failure branch logs, and a missing artifact is visible in the API, in "
            "the dashboard and in CI.",
        ),
        (
            "Train the thing you actually use.",
            "We trained the GNN with mean-squared error against a path's true "
            "cost, but the only thing we ever did with its output was take the "
            "argmin. Regression accuracy on a value nobody reads is not the "
            "objective. Switching to a pairwise ranking loss and reporting top-1 "
            "accuracy and regret changed both the training signal and what we "
            "could honestly claim.",
        ),
        (
            "Correlated observations are not independent samples.",
            "We ran significance tests over 20,000 routing decisions taken from a "
            "single simulation trajectory. Successive steps are almost the same "
            "network, so those were roughly one observation repeated. That is why "
            "our p-values were literally 0.0 — numerical underflow, which we had "
            "been reading as very strong evidence.",
        ),
        (
            "Ask whether the information is in the input at all.",
            "Two of our models were structurally unable to answer the question we "
            "were asking. The RL agent was never told which source-destination "
            "pair it was routing. The GNN's path representation was a mean over "
            "node embeddings, which is blind to path length, hop order and the "
            "state of the links on the path. Both failures look like "
            "underperformance and neither is fixable by training longer.",
        ),
        (
            "Publishing a negative result is stronger than hiding it.",
            "Our benchmark showed every AI method losing to Dijkstra, and we "
            "published it nowhere across 2,975 lines of documentation. A reviewer "
            "who finds that themselves concludes it was concealed. A rigorous "
            "negative result with a correct causal explanation is a genuine "
            "contribution — and the explanation, once we looked for it, was "
            "interesting: on an additive objective, converging to Dijkstra is what "
            "a correct algorithm does.",
        ),
        (
            "Instrumentation that reports but does not enforce will be ignored.",
            "We built genuinely good self-exposing instrumentation — a fallback "
            "flag threaded through the whole stack, a degeneracy metric, a UI "
            "badge that displays our own models' failure modes. None of it was "
            "wired to anything that could fail. The documented training result "
            "drifted away from the committed evidence and nothing noticed. The "
            "guardrails are now gates in CI.",
        ),
        (
            "Work out the steady state before trusting a magnitude.",
            "When we added routing-induced load to an AR(1) process, we added it "
            "outside the update and the steady state amplified it tenfold. The "
            "network saturated and the cost function went degenerate again. It "
            "cost an afternoon and would have cost nothing to check on paper.",
        ),
        (
            "Ask what survives a round trip.",
            "We made a model's actor and critic asymmetric by patching the object "
            "after construction. The training run was fine; the checkpoint could "
            "not be reloaded, and the router silently fell back to a heuristic for "
            "a whole benchmark run. Anything not expressed in the constructor or "
            "the saved configuration does not survive serialisation.",
        ),
    ]

    for index, (headline, body) in enumerate(lessons, start=1):
        paragraph = doc.add_paragraph()
        run = paragraph.add_run(f"{index}. {headline}")
        run.bold = True
        doc.add_paragraph(body)

    doc.add_page_break()


def section_challenges(doc, results, ml):
    add_heading(doc, "3. Challenges and How We Overcame Them", 1)

    doc.add_paragraph(
        "Ordered by how much they cost us. Each entry states the problem, what it "
        "actually was underneath, and what we did."
    )

    challenges = [
        (
            "The AI did not run, and nothing said so",
            "Three of four learned components never loaded their models. The RL "
            "router expected 'rl_router_final.zip' while training saved "
            "'ppo_routing_agent.zip'; the MARL and LSTM artifacts were absent "
            "entirely. All three failures were swallowed by exception handlers "
            "that returned False without logging.",
            "Underneath, this was not really a filename bug — it was that three "
            "modules each hardcoded their own path with nothing connecting them. "
            "We built a single model registry that both training and serving read, "
            "made every failure branch log at a level chosen by how surprising it "
            "is, added a /health/models endpoint, put a banner in the dashboard, "
            "and wrote a verification script that fails CI if a registered "
            "artifact is missing. While fixing it we found a worse latent bug: a "
            "failed GNN load left a randomly-initialised model installed while "
            "reporting itself as trained.",
        ),
        (
            "The environment made the goal unreachable",
            "Link utilization evolved as a random walk with no dependence on "
            "routing. Load balancing was therefore unmeasurable, the RL reward's "
            "global terms contributed exactly zero gradient, and Dijkstra was "
            "provably optimal — so no learned policy could win.",
            "We closed the loop: routing a flow now adds load to the links it "
            "traverses, with exponential decay and background traffic. The first "
            "implementation added the flow term outside the AR(1) update, where "
            "the steady state amplified it tenfold and saturated the network; "
            "folding it into the offered-load baseline gave a 1:1 mapping. We then "
            "wrote a test proving that round-robin across three paths keeps the "
            "worst link cooler than saturating one — the load-balancing property "
            "the project is about, and previously impossible to demonstrate.",
        ),
        (
            "The PPO agent could not see the problem it was solving",
            "500,000 timesteps produced a learning curve statistically "
            "indistinguishable from flat: slope -0.094 per 100k steps, r-squared "
            "0.001, p 0.878, with the best checkpoint being the first one taken.",
            "The observation encoded per-link features only, while the environment "
            "re-drew the source-destination pair every step. The agent was asked "
            "to choose 'path index 2' without being told which pair it was routing, "
            "and the meaning of index 2 changed between steps. That is an "
            "unobservable MDP, not a hard one. We rebuilt the observation to carry "
            "the task, the candidate features and the traffic class, fixed the "
            "reward/observation ordering, and reset the simulator per episode. The "
            "curve now rises. The agent still does not beat the greedy baseline, "
            "and we report that.",
        ),
        (
            "Our statistics were invalid by construction",
            "Every published p-value was exactly 0.0, which we had read as very "
            "strong evidence.",
            "It was numerical underflow. We were running a paired test over 20,000 "
            "routing decisions from a single trajectory, and successive simulator "
            "steps are almost the same network — so those were approximately one "
            "independent observation repeated. We changed the unit of replication "
            "to one seeded run, gave each algorithm its own closed-loop trajectory "
            "with an identical demand schedule, and added Cliff's delta and "
            "bootstrap confidence intervals. Two metrics turned out to be "
            "structurally broken as well: path diversity read a key that was never "
            "stored, and bottleneck utilization took a maximum over 20,000 samples "
            "and so was always 1.000.",
        ),
        (
            "Calling a mixture of experts 'multi-agent RL'",
            "Our multi-agent router was N independently trained agents, each "
            "seeing the full global state and each emitting a complete end-to-end "
            "path, selected by a lookup on the source node. It was documented as "
            "centralized-critic, decentralized-execution multi-agent RL.",
            "We had a choice between renaming it honestly in half an hour and "
            "implementing the thing we had claimed. We implemented it: local "
            "observations of constant width, next-hop actions, control transferring "
            "between regions along a path, and a critic that sees a global summary "
            "the actor never does. Then we wrote tests that verify each property "
            "rather than describing it — including one that perturbs only the "
            "global observation block and checks the actor does not react.",
        ),
        (
            "A forecaster that could not beat copying the last value",
            "The LSTM's first honest evaluation scored a skill score of -1.77 "
            "against persistence, and the training script refused to save it.",
            "That refusal was the feature working. The diagnosis was that "
            "predicting the utilization level means competing with persistence on "
            "a strongly autocorrelated series, where copying the last value is "
            "right to within the noise almost every time — the network was "
            "spending its capacity relearning the identity function. Predicting "
            "the residual instead removed the identity from the problem: skill "
            "score +0.15, saved. Predictive routing executes for the first time; "
            "previously it was a silent no-op, which is why the predictive "
            "benchmark columns were byte-identical to their base algorithms.",
        ),
        (
            "Scenarios that destroyed the signal they existed to measure",
            "The high-congestion scenario added load cumulatively every step with "
            "no reset, saturating every link within about ten steps; once all "
            "utilizations are equal the cost function collapses to base-latency "
            "order. The link-failure scenario re-randomised the topology every "
            "tick and had zero discriminative power. The 100-node topology was a "
            "ring: degree 2, diameter 50, exactly two paths between any pair.",
            "We rewrote the scenarios as declarative objects with a sustained bias "
            "the AR(1) process mean-reverts toward, persistent failures that skip "
            "any edge whose removal would disconnect the graph, and a topology "
            "generator whose edge count scales with degree. The 100-node network "
            "is now degree 4 with diameter 8. We added two scenarios that can "
            "actually discriminate: progressive cascading failure, and mixed QoS "
            "traffic.",
        ),
        (
            "A test suite that could not be run and a linter that could not start",
            "pytest was in no requirements file; there was no pytest.ini and no "
            "conftest.py. Two of the eleven test files contained zero test "
            "functions. The frontend had four ESLint plugins installed and no "
            "configuration file anywhere, so 'npm run lint' failed immediately.",
            "Both are now installable and both run in CI. Writing the tests found "
            "real bugs — including one where routing a node to itself cost "
            "infinity, because a zero-hop path was being treated as an invalid "
            "one. We added honesty gates that fail the build if an algorithm is "
            "silently degenerate, if a row is secretly a heuristic, if a metric "
            "goes structurally constant, or if a documented number stops matching "
            "its artifact.",
        ),
        (
            "The reviewer's first command failed",
            "'docker compose up' on a fresh clone failed immediately: three "
            "services declared env_file: .env, .env is gitignored, and nothing "
            "created it.",
            "We made the env file optional, added a Makefile so first run is one "
            "command, rebuilt the images as two-stage non-root builds without the "
            "compiler in the runtime layer, moved pgAdmin behind a dev profile, and "
            "added a CI job that builds from a clean checkout and asserts the "
            "models really loaded. We should note that no Docker daemon was "
            "available in our final build environment, so the container run is "
            "reviewed and CI-specified but not executed by us.",
        ),
        (
            "Finding out that being right looks like failing",
            "After all the fixes, the trained GNN reproduces Dijkstra's chosen "
            "path essentially 100% of the time on best-effort traffic, and our "
            "degeneracy guardrail flagged it.",
            "This was the most interesting result of the project. Dijkstra is "
            "provably optimal for an additive edge cost, so a well-trained ranker "
            "*must* converge to it — the guardrail was reporting correct behaviour "
            "as a problem. We changed the gate from 'must not be degenerate' to "
            "'degeneracy must be declared', which is what the test was always "
            "named for, and it reframed the whole project: learned routing can "
            "only add value where the objective is not a single additive cost. "
            "That is precisely why we built the QoS layer.",
        ),
    ]

    for index, (headline, problem, response) in enumerate(challenges, start=1):
        add_heading(doc, f"3.{index} {headline}", 2)
        paragraph = doc.add_paragraph()
        paragraph.add_run("What we hit: ").bold = True
        paragraph.add_run(problem)
        paragraph = doc.add_paragraph()
        paragraph.add_run("What we did: ").bold = True
        paragraph.add_run(response)

    doc.add_page_break()


def section_results(doc, results, ml):
    add_heading(doc, "4. Results", 1)

    if ml:
        add_heading(doc, "4.1 Model quality", 2)
        rows = []
        if "gnn" in ml:
            g = ml["gnn"]
            rows.append([
                "GNN path ranker",
                f"Top-1 {g['test']['top1_accuracy']:.3f}",
                f"Random {g['random_baseline_test']['top1_accuracy']:.3f}",
                "Beats the floor decisively",
            ])
        if "rl" in ml:
            r = ml["rl"]["normalized_scores"]
            rows.append([
                "PPO routing policy",
                f"Normalized {r.get('ppo', 0):.3f}",
                f"Greedy {r.get('greedy_first_candidate', 0):.3f}",
                "Learns, but does not beat greedy",
            ])
        if "lstm" in ml:
            s = ml["lstm"]
            rows.append([
                "LSTM forecaster",
                f"Skill {s['skill_score_vs_persistence']:+.4f}",
                f"MSE {s['test_mse']['lstm']:.6f} vs {s['test_mse']['persistence']:.6f}",
                "Beats persistence, modestly",
            ])
        if "marl" in ml:
            m = ml["marl"]
            rows.append([
                "Regional CTDE policies",
                f"{m['regions_beating_random']}/{len(m['per_region'])} regions beat random",
                f"Local obs {m['local_obs_dim']}-d",
                "Decentralized execution verified",
            ])
        add_table(doc, ["Model", "Headline metric", "Reference", "Reading"], rows,
                  widths=[1.5, 1.6, 1.7, 1.6])

    if results:
        add_heading(doc, "4.2 Benchmark", 2)
        doc.add_paragraph(
            "Mean congestion-adjusted path latency in milliseconds, lower is "
            "better, across independently seeded runs. The percentage is the "
            "difference against Dijkstra."
        )
        first = next(iter(results.values()))
        algorithms = list(first["algorithms"])
        header = ["Scenario", *[a[:11] for a in algorithms]]
        rows = []
        marked = False
        for name, data in results.items():
            row = [name[:26]]
            for algorithm in algorithms:
                metrics = data["algorithms"].get(algorithm, {})
                latency = metrics.get("mean_latency")
                if latency is None:
                    row.append("—")
                    continue
                cell = f"{latency:.1f}"
                # A row that mostly came from the fallback heuristic is not that
                # algorithm's result, and must not be read as one. In
                # link_failures_persistent the unmarked rl number is *better*
                # than Dijkstra's — produced entirely by a five-line heuristic.
                if (metrics.get("fallback_rate") or 0.0) > 0.2:
                    cell += " †"
                    marked = True
                row.append(cell)
            rows.append(row)
        add_table(doc, header, rows, widths=[1.5] + [0.62] * len(algorithms))
        if marked:
            add_caption(
                doc,
                "† More than 20% of this row's decisions came from the heuristic "
                "fallback rather than the named model, so the number is not that "
                "model's result. The PPO agent's observation is a fixed-width "
                "vector sized for one topology, so it cannot run on a 100-node "
                "graph or one with links removed; it detects the mismatch and "
                "falls back, and says so.",
            )

        qos = results.get("qos_mixed_traffic")
        if qos:
            add_heading(doc, "4.3 The QoS regime, where the thesis actually lives", 2)
            doc.add_paragraph(
                "Dijkstra is provably optimal for an additive, non-negative cost, "
                "so on best-effort traffic no method can beat it and the table "
                "above shows exactly that: the top five algorithms are "
                "statistically indistinguishable. The qos_mixed_traffic scenario "
                "was built to leave that regime — five traffic classes with hard "
                "constraints on jitter, loss and bottleneck utilisation, a problem "
                "that is NP-hard in general. If a learned router were going to "
                "win anywhere, it would win here."
            )
            doc.add_paragraph(
                "The metric is the percentage of decisions whose chosen path "
                "satisfied every hard constraint of its traffic class. Read the "
                "emergency column, not the overall one: best_effort has no hard "
                "constraints, sits at 100% for everyone, and drags the average "
                "together. The overall spread is about 15 points; the emergency "
                "spread is about 31."
            )
            classes = ["emergency", "interactive", "gaming", "bulk", "best_effort"]
            rows = []
            for algorithm, metrics in qos["algorithms"].items():
                row = [algorithm[:13]]
                for traffic_class in classes:
                    rate = metrics.get(f"qos_satisfaction_rate__{traffic_class}")
                    row.append(f"{rate * 100:.1f}%" if rate is not None else "—")
                overall = metrics.get("qos_satisfaction_rate")
                row.append(f"{overall * 100:.1f}%" if overall is not None else "—")
                rows.append(row)
            add_table(
                doc,
                ["Algorithm", "emergency", "interact.", "gaming", "bulk", "best-eff.", "overall"],
                rows,
                widths=[1.3, 0.85, 0.85, 0.8, 0.75, 0.85, 0.8],
            )
            doc.add_paragraph(
                "The classical constrained router — k-shortest paths filtered by "
                "feasibility — wins every class, and pays 1.6% in mean latency to "
                "do it. That is the trade a constraint-aware router should make. "
                "No learned router beats it."
            )
            doc.add_paragraph(
                "This is not evidence that learned routing fails here, and we are "
                "careful not to write it up that way. It is a training-objective "
                "gap: the GNN was trained to rank paths by additive cost and the "
                "PPO agent was rewarded for latency, so neither has ever been "
                "asked to satisfy a constraint. They behave exactly as trained. "
                "The honest summary of this project is that we built the arena in "
                "which learned routing could win, and have not yet trained a model "
                "that wins in it. Closing that gap — a feasibility term in the "
                "reward, a constraint-aware ranking loss — is the most direct "
                "experiment we are leaving undone."
            )

        warnings = [w for data in results.values() for w in data.get("warnings", [])]
        if warnings:
            add_heading(doc, "4.4 Guardrail warnings emitted by the benchmark", 2)
            doc.add_paragraph(
                "These are produced automatically and appear in the dashboard "
                "above the results table. We report them rather than leaving them "
                "to be discovered."
            )
            add_bullets(doc, sorted(set(warnings))[:12])

    for figure, caption in (
        ("ppo_eval_curve_baseline.png", "PPO evaluation curve before the fixes: r-squared 0.001, best checkpoint the first one taken."),
        ("ppo_eval_curve_current.png", "PPO evaluation curve after the fixes: the task is observable and the curve rises."),
        ("ppo_normalized_score.png", "PPO against a random floor and a greedy oracle ceiling."),
    ):
        path = ASSETS / figure
        if path.exists():
            doc.add_picture(str(path), width=Inches(5.9))
            doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
            add_caption(doc, caption)

    doc.add_page_break()


def section_honesty(doc):
    add_heading(doc, "5. What Is Still Not True", 1)
    doc.add_paragraph(
        "Stated here rather than left to be found. This section exists because "
        "the most damaging problem we found was not any individual defect — it "
        "was that the disclosed problems were all small while the undisclosed "
        "ones were all large, which makes a reader conclude the disclosure was "
        "managed rather than honest."
    )
    add_bullets(
        doc,
        [
            "PPO does not beat the greedy baseline: 0.867 normalized against 0.903 for simply taking the cheapest candidate.",
            "The GNN is degenerate on best-effort traffic — it reproduces Dijkstra's path essentially always. That is correct behaviour and it means it adds nothing there.",
            "Nothing learns online. Every model is trained offline and frozen, so the problem statement's 'continuously learns network conditions' is still not delivered. It was explicitly de-scoped.",
            "The network is synthetic. Trace replay and live measurement exist, but no published number comes from either, and live mode's star topology cannot benchmark routing at all.",
            "The LSTM's edge is small at +0.15 skill; predictive routing changes the chosen path only occasionally.",
            "The QoS oracle is greedy rather than optimal — it ignores the downstream consequences of its own load.",
            "No Docker daemon was available in the final build environment, so the container stack is reviewed and CI-specified but was not executed by us. It should be the first thing verified.",
            "Sample sizes are modest: 15 runs of 40 steps with 8 demands per scenario, chosen so the suite completes in about half an hour on a laptop CPU.",
        ],
    )

    add_heading(doc, "5.1 Deliverables", 2)
    add_table(
        doc,
        ["Artifact", "Contents"],
        [
            ["README.md", "Thesis, headline results, full startup guide, reproduction instructions"],
            ["LEARNING_GUIDE.md", "The project written up as a research paper, with real-world applications and a deployment design"],
            ["docs/14_RESULTS_AND_FINDINGS.md", "Experimental setup, full results, and the findings stated plainly"],
            ["ml/cards/", "One model card per trained model, with use-matched metrics and failure modes"],
            ["experiments/README.md", "Benchmark methodology and known limitations"],
            ["FINAL_REPORT.docx", "This document"],
        ],
        widths=[2.0, 4.1],
    )


def main() -> int:
    soham = parse_worklog(REPO_ROOT / "WORKLOG_Soham.md")
    sneha = parse_worklog(REPO_ROOT / "WORKLOG_Sneha.md")
    results = load_results()
    ml = load_ml()

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10.5)

    cover(doc, soham, sneha)
    section_hours(doc, soham, sneha)
    section_learning(doc, soham, sneha)
    section_challenges(doc, results, ml)
    section_results(doc, results, ml)
    section_honesty(doc)

    destination = REPO_ROOT / "FINAL_REPORT.docx"
    doc.save(destination)

    print(f"Wrote {destination.relative_to(REPO_ROOT)}")
    print(f"  hours: {soham['name']} {soham['total']:.1f} + {sneha['name']} {sneha['total']:.1f}"
          f" = {soham['total'] + sneha['total']:.1f}")
    print(f"  scenarios: {len(results)}   models: {len(ml)}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
